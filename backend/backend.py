from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from dataclasses import asdict
from pydantic import BaseModel
from typing import Dict     
import tempfile
import shutil
import os
import json
import requests

from legal_analyzer import LegalDocumentAnalyzer
from contract_generator import ContractTemplateGenerator

# =========================
# FastAPI app
# =========================
app = FastAPI()
analyzer = LegalDocumentAnalyzer(verbose=False)
contract_generator = ContractTemplateGenerator()

# Allow React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173"],  # React/Vite dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# Pydantic Models
# =========================
class ChatRequest(BaseModel):
    question: str
    document_text: str

class ContractGenerationRequest(BaseModel):
    contract_type: str
    user_data: Dict[str, str]
    format_type: str = "docx"

# =========================
# Helper function
# =========================
def save_temp_file(upload_file: UploadFile) -> str:
    ext = os.path.splitext(upload_file.filename)[1].lower()
    if ext not in [".pdf", ".docx", ".txt"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        shutil.copyfileobj(upload_file.file, tmp)
        return tmp.name

def cleanup_temp_file(file_path: str):
    """Safely remove temporary file"""
    try:
        if os.path.exists(file_path):
            os.unlink(file_path)
    except Exception as e:
        print(f"Error cleaning up temp file: {e}")

def cleanup_temp_dir(dir_path: str):
    """Safely remove temporary directory"""
    try:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
    except Exception as e:
        print(f"Error cleaning up temp directory: {e}")

# =========================
# Routes
# =========================
@app.get("/")
async def root():
    return {
        "message": "Legal Document Analyzer API",
        "endpoints": {
            "/analyze": "POST - Analyze a single document",
            "/compare": "POST - Compare two documents",
            "/chat": "POST - Chat with document using AI",
            "/contract-templates": "GET - List available contract templates",
            "/generate-contract": "POST - Generate contract from built-in templates",
            "/generate-contract-from-custom-template": "POST - Generate contract from custom template"
        }
    }

@app.post("/analyze")
async def analyze_document(file: UploadFile = File(...)):
    file_path = save_temp_file(file)
    try:
        results = analyzer.analyze(file_path)
        return asdict(results)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cleanup_temp_file(file_path)

@app.post("/compare")
async def compare_documents(file1: UploadFile = File(...), file2: UploadFile = File(...)):
    path1, path2 = save_temp_file(file1), save_temp_file(file2)
    try:
        results = analyzer.compare_documents(path1, path2)
        return asdict(results)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        cleanup_temp_file(path1)
        cleanup_temp_file(path2)

@app.post("/chat")
async def chat_with_document(request: ChatRequest):
    try:
        prompt = f"""You are a legal document assistant. Based on the following document, answer the user's question accurately and concisely.

Document Content:
{request.document_text[:4000]}

User Question: {request.question}

Answer:"""

        ollama_response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3.2",
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.7,
                    "top_p": 0.9,
                }
            },
            timeout=60
        )

        if not ollama_response.ok:
            raise HTTPException(
                status_code=500, 
                detail=f"Ollama API error: {ollama_response.text}"
            )

        ollama_data = ollama_response.json()
        answer = ollama_data.get("response", "").strip()

        if not answer:
            raise HTTPException(
                status_code=500,
                detail="No response from Ollama. Make sure the model is loaded."
            )

        return {"answer": answer}

    except requests.exceptions.ConnectionError:
        raise HTTPException(
            status_code=503,
            detail="Cannot connect to Ollama. Make sure Ollama is running on localhost:11434"
        )
    except requests.exceptions.Timeout:
        raise HTTPException(
            status_code=504,
            detail="Ollama request timed out. The question might be too complex."
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# =========================
# Contract Generation Routes
# =========================
@app.get("/contract-templates")
async def get_contract_templates():
    """Get available contract templates with their required fields"""
    try:
        templates = contract_generator.get_available_templates()
        
        template_details = {}
        for template_type, description in templates.items():
            template_details[template_type] = {
                "description": description,
                "required_fields": contract_generator.get_required_fields(template_type)
            }
        
        return {
            "success": True,
            "templates": template_details
        }
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-contract")
async def generate_builtin_contract(request: ContractGenerationRequest):
    """
    Generate a contract from built-in templates (NDA, Service Agreement).
    
    Request body (JSON):
    {
        "contract_type": "nda" | "service_agreement",
        "user_data": {
            "field_name": "field_value",
            ...
        },
        "format_type": "docx" | "txt"
    }
    """
    temp_dir = None
    generated_path = None
    
    try:
        # Validate contract type
        available_templates = contract_generator.get_available_templates()
        if request.contract_type not in available_templates:
            raise HTTPException(
                status_code=400, 
                detail=f"Invalid contract type. Available: {list(available_templates.keys())}"
            )
        
        # Validate format type
        if request.format_type not in ["docx", "txt"]:
            raise HTTPException(
                status_code=400,
                detail="format_type must be 'docx' or 'txt'"
            )
        
        # Create temp directory for output
        temp_dir = tempfile.mkdtemp()
        import time
        output_filename = f"{request.contract_type}_{int(time.time())}.{request.format_type}"
        output_path = os.path.join(temp_dir, output_filename)
        
        # Generate contract
        generated_path = contract_generator.generate_contract(
            contract_type=request.contract_type,
            user_data=request.user_data,
            output_path=output_path,
            format_type=request.format_type
        )
        
        # Determine media type
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            if request.format_type == "docx"
            else "text/plain"
        )
        
        # Return file for download
        return FileResponse(
            generated_path,
            media_type=media_type,
            filename=output_filename,
            background=lambda: cleanup_temp_dir(temp_dir)
        )
        
    except ValueError as e:
        if temp_dir:
            cleanup_temp_dir(temp_dir)
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        if temp_dir:
            cleanup_temp_dir(temp_dir)
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate contract: {str(e)}")

@app.post("/generate-contract-from-custom-template")
async def generate_custom_contract(
    template_file: UploadFile = File(...), 
    fields_json: str = Form(...)
):
    """
    Generate a contract from a user-uploaded DOCX or TXT template.
    The template should contain placeholders like {{party_name}}.
    
    Form data:
    - template_file: DOCX or TXT file with placeholders
    - fields_json: JSON string with field values, e.g. {"party_name": "John Doe"}
    """
    template_path = None
    output_path = None
    temp_dir = None
    
    try:
        # Parse fields JSON
        try:
            fields = json.loads(fields_json)
        except json.JSONDecodeError:
            raise HTTPException(
                status_code=400, 
                detail="Invalid fields_json. Must be valid JSON string."
            )
        
        # Validate file extension
        ext = os.path.splitext(template_file.filename)[1].lower()
        if ext not in [".docx", ".txt"]:
            raise HTTPException(
                status_code=400, 
                detail="Template must be DOCX or TXT file"
            )
        
        # Save uploaded template temporarily
        temp_dir = tempfile.mkdtemp()
        template_path = os.path.join(temp_dir, f"template{ext}")
        
        with open(template_path, "wb") as buffer:
            shutil.copyfileobj(template_file.file, buffer)
        
        # Generate output path
        output_path = os.path.join(temp_dir, f"generated_contract{ext}")
        
        # Fill template with user data
        if ext == ".docx":
            from docx import Document
            doc = Document(template_path)
            
            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for key, value in fields.items():
                    placeholder = f"{{{{{key}}}}}"
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in fields.items():
                                placeholder = f"{{{{{key}}}}}"
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))
            
            doc.save(output_path)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        else:  # .txt
            with open(template_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            for key, value in fields.items():
                content = content.replace(f"{{{{{key}}}}}", str(value))
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            
            media_type = "text/plain"
        
        # Return the generated file
        return FileResponse(
            output_path,
            media_type=media_type,
            filename=f"generated_contract{ext}",
            background=lambda: cleanup_temp_dir(temp_dir)
        )
        
    except json.JSONDecodeError:
        if temp_dir:
            cleanup_temp_dir(temp_dir)
        raise HTTPException(status_code=400, detail="Invalid fields_json format")
    except Exception as e:
        if temp_dir:
            cleanup_temp_dir(temp_dir)
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate contract: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "analyzer": "ready",
        "contract_generator": "ready"
    }
