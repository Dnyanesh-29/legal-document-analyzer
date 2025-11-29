# from docx import Document
# from docx.shared import Pt, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from datetime import datetime
# from typing import Dict
# import os
# import re

# class ContractTemplateGenerator:
#     """
#     Enhanced contract generator with multiple template types.
#     Supports filling placeholders with user data and generating downloadable contracts.
#     """
    
#     def __init__(self):
#         self.templates = {
#             "nda": self._get_nda_template(),
#             "service_agreement": self._get_service_agreement_template(),
#         }
    
#     def _get_nda_template(self) -> str:
#         """Non-Disclosure Agreement (NDA) Template"""
#         return """
# NON-DISCLOSURE AGREEMENT

# This Non-Disclosure Agreement ("Agreement") is entered into as of {{effective_date}}, by and between:

# DISCLOSING PARTY:
# {{disclosing_party_name}}
# {{disclosing_party_address}}
# {{disclosing_party_email}}

# and

# RECEIVING PARTY:
# {{receiving_party_name}}
# {{receiving_party_address}}
# {{receiving_party_email}}

# (collectively referred to as the "Parties")

# WHEREAS, the Disclosing Party possesses certain confidential and proprietary information related to {{business_purpose}};

# WHEREAS, the Receiving Party desires to receive such confidential information for the purpose of {{purpose_of_disclosure}};

# NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, the Parties agree as follows:

# 1. DEFINITION OF CONFIDENTIAL INFORMATION

# "Confidential Information" means any and all information disclosed by the Disclosing Party to the Receiving Party, whether orally, in writing, or in any other form, including but not limited to:
#    a) Technical data, trade secrets, know-how, research, product plans, products, services
#    b) Software, algorithms, source code, object code
#    c) Customer lists, supplier information, business strategies
#    d) Financial information, forecasts, budgets
#    e) Marketing plans, sales data
#    f) Any other information marked as "Confidential" or that would reasonably be considered confidential

# 2. OBLIGATIONS OF RECEIVING PARTY

# The Receiving Party agrees to:
#    a) Hold and maintain the Confidential Information in strict confidence
#    b) Not disclose the Confidential Information to any third parties without prior written consent
#    c) Use the Confidential Information solely for the purpose of {{purpose_of_disclosure}}
#    d) Protect the Confidential Information with the same degree of care used to protect its own confidential information, but in no case less than reasonable care
#    e) Limit access to the Confidential Information to employees and contractors who have a legitimate need to know

# 3. EXCLUSIONS FROM CONFIDENTIAL INFORMATION

# Confidential Information shall not include information that:
#    a) Is or becomes publicly available through no breach of this Agreement
#    b) Was rightfully in the Receiving Party's possession prior to disclosure
#    c) Is independently developed by the Receiving Party without use of the Confidential Information
#    d) Is rightfully received from a third party without breach of any confidentiality obligation

# 4. TERM AND TERMINATION

# 4.1 Term
# This Agreement shall commence on the Effective Date and continue for a period of {{term_years}} years, unless terminated earlier by either Party with {{notice_period}} days written notice.

# 4.2 Survival
# The obligations of confidentiality shall survive termination for a period of {{survival_years}} years. After such period, the obligations under this Agreement shall cease and terminate.

# 4.3 Right to Terminate
# Either Party may terminate this Agreement at any time with {{notice_period}} days prior written notice to the other Party.

# 5. RETURN OF MATERIALS

# Upon termination of this Agreement or upon request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and certify such destruction in writing within thirty (30) days.

# 6. NO LICENSE

# Nothing in this Agreement grants the Receiving Party any license or right to the Confidential Information except as expressly stated herein.

# 7. INDEMNIFICATION

# The Receiving Party agrees to indemnify, defend, and hold harmless the Disclosing Party from and against any and all claims, damages, liabilities, costs, and expenses (including reasonable attorneys' fees) arising from or related to:
#    a) Any breach of this Agreement by the Receiving Party
#    b) Any unauthorized use or disclosure of Confidential Information by the Receiving Party
#    c) Any negligent or willful misconduct by the Receiving Party

# 8. LIMITATION OF LIABILITY

# The liability of each Party under this Agreement shall not exceed the actual damages suffered, up to a maximum of $100,000. In no event shall either Party be liable for indirect, incidental, special, consequential, or punitive damages.

# 9. DISPUTE RESOLUTION

# 9.1 Good Faith Negotiation
# In the event of any dispute arising out of or relating to this Agreement, the Parties shall first attempt to resolve the dispute through good faith negotiations.

# 9.2 Mediation
# If the dispute cannot be resolved through negotiation within thirty (30) days, the Parties agree to submit the dispute to mediation before a mutually agreed-upon mediator.

# 9.3 Arbitration
# If mediation is unsuccessful, any remaining disputes shall be resolved through binding arbitration in accordance with the rules of the American Arbitration Association. The arbitration shall take place in {{governing_state}}.

# 10. REMEDIES

# The Receiving Party acknowledges that breach of this Agreement may cause irreparable harm for which monetary damages would be inadequate. The Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity.

# 11. GOVERNING LAW AND JURISDICTION

# This Agreement shall be governed by and construed in accordance with the laws of the State of {{governing_state}}, without regard to its conflict of law provisions. The Parties consent to the exclusive jurisdiction of the courts located in {{governing_state}} for any legal proceedings.

# 12. ENTIRE AGREEMENT

# This Agreement constitutes the entire agreement between the Parties concerning the subject matter hereof and supersedes all prior agreements and understandings, whether written or oral.

# 13. AMENDMENTS AND MODIFICATIONS

# This Agreement may only be amended or modified by a written document signed by both Parties. No waiver of any provision shall be deemed a waiver of any other provision.

# 14. SEVERABILITY

# If any provision of this Agreement is found to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.

# 15. ASSIGNMENT

# Neither Party may assign this Agreement without the prior written consent of the other Party, except in connection with a merger, acquisition, or sale of substantially all assets.

# IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

# DISCLOSING PARTY:

# _________________________________
# {{disclosing_party_name}}
# {{disclosing_party_title}}
# Date: {{signature_date}}


# RECEIVING PARTY:

# _________________________________
# {{receiving_party_name}}
# {{receiving_party_title}}
# Date: {{signature_date}}
# """

#     def _get_service_agreement_template(self) -> str:
#         """Service Agreement Template"""
#         return """
# SERVICE AGREEMENT

# This Service Agreement ("Agreement") is entered into as of {{effective_date}}, by and between:

# SERVICE PROVIDER:
# {{provider_name}}
# {{provider_address}}
# {{provider_email}}
# {{provider_phone}}

# and

# CLIENT:
# {{client_name}}
# {{client_address}}
# {{client_email}}
# {{client_phone}}

# (collectively referred to as the "Parties")

# WHEREAS, the Service Provider is engaged in the business of providing {{service_type}} services;

# WHEREAS, the Client desires to engage the Service Provider to perform certain services as described herein;

# NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, the Parties agree as follows:

# 1. SERVICES

# The Service Provider agrees to provide the following services ("Services"):

# {{service_description}}

# The Services shall be performed in accordance with industry standards and professional practices.

# 2. TERM AND TERMINATION

# 2.1 Initial Term
# This Agreement shall commence on {{start_date}} and continue until {{end_date}} ("Initial Term"), unless terminated earlier in accordance with this Agreement.

# 2.2 Renewal
# This Agreement may be renewed for additional {{renewal_period}} periods upon mutual written agreement of the Parties at least {{renewal_notice_days}} days prior to expiration.

# 2.3 Termination for Convenience
# Either Party may terminate this Agreement with {{termination_notice}} days written notice to the other Party.

# 2.4 Termination for Cause
# Either Party may terminate this Agreement immediately upon written notice if the other Party:
#    a) Materially breaches this Agreement and fails to cure within {{cure_period}} days of written notice
#    b) Becomes insolvent or files for bankruptcy
#    c) Engages in fraudulent or illegal conduct

# 2.5 Effect of Termination
# Upon termination, the Client shall pay for all Services performed up to the termination date and any reasonable expenses incurred. All obligations that by their nature should survive termination shall remain in effect.

# 3. COMPENSATION

# 3.1 Service Fees
# The Client shall pay the Service Provider as follows:

# Payment Structure: {{payment_structure}}
# Total Amount: {{total_amount}}
# Payment Schedule: {{payment_schedule}}

# 3.2 Additional Expenses
# The Service Provider shall be reimbursed for pre-approved, reasonable expenses incurred in the performance of Services, including travel expenses, materials and supplies, and third-party services.

# All expenses must be documented with receipts and approved by the Client in writing prior to incurrence.

# 3.3 Payment Terms
# Invoices shall be submitted {{invoice_frequency}} and are due within {{payment_terms}} days of receipt. Late payments shall accrue interest at {{late_fee_rate}}% per month.

# 4. DELIVERABLES

# The Service Provider shall deliver the following to the Client:

# {{deliverables}}

# All deliverables shall be delivered by {{delivery_date}}, subject to any agreed-upon extensions.

# 5. INDEPENDENT CONTRACTOR

# The Service Provider is an independent contractor and not an employee of the Client. The Service Provider shall be responsible for all taxes, insurance, and other obligations related to their status as an independent contractor.

# 6. INTELLECTUAL PROPERTY

# 6.1 Client Materials
# All materials, information, and intellectual property provided by the Client remain the property of the Client.

# 6.2 Work Product
# {{ip_ownership_clause}}

# 6.3 Pre-Existing Materials
# Any pre-existing materials, tools, or intellectual property owned by the Service Provider prior to this Agreement shall remain the property of the Service Provider.

# 7. CONFIDENTIALITY

# Both Parties agree to maintain the confidentiality of any proprietary or confidential information disclosed during the term of this Agreement. This obligation shall survive termination for a period of {{confidentiality_period}} years. After such period, the confidentiality obligations shall cease and terminate.

# 8. WARRANTIES

# The Service Provider warrants that:
#    a) Services will be performed in a professional and workmanlike manner
#    b) Services will comply with all applicable laws and regulations
#    c) Service Provider has the right and authority to enter into this Agreement
#    d) Services will not infringe upon any third-party intellectual property rights

# 9. INDEMNIFICATION

# 9.1 Service Provider Indemnification
# The Service Provider agrees to indemnify, defend, and hold harmless the Client from any claims, damages, liabilities, costs, and expenses (including reasonable attorneys' fees) arising from:
#    a) Any breach of this Agreement by the Service Provider
#    b) Any negligent or willful misconduct by the Service Provider
#    c) Any infringement of third-party intellectual property rights by the Services

# 9.2 Client Indemnification
# The Client agrees to indemnify, defend, and hold harmless the Service Provider from any claims arising from:
#    a) The Client's use of the Services beyond the scope of this Agreement
#    b) Any materials or information provided by the Client
#    c) Any breach of this Agreement by the Client

# 10. LIMITATION OF LIABILITY

# 10.1 Cap on Liability
# The Service Provider's total liability under this Agreement shall not exceed {{liability_cap}}.

# 10.2 Exclusion of Damages
# IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOST PROFITS, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.

# 11. DISPUTE RESOLUTION

# 11.1 Good Faith Negotiation
# In the event of any dispute arising out of or relating to this Agreement, the Parties shall first attempt to resolve the dispute through good faith negotiations for a period of thirty (30) days.

# 11.2 Mediation
# If the dispute cannot be resolved through negotiation, the Parties agree to submit the dispute to mediation before a mutually agreed-upon mediator.

# 11.3 Arbitration or Litigation
# If mediation is unsuccessful, the dispute shall be resolved through:
# {{dispute_resolution_method}}

# The arbitration or litigation shall take place in {{governing_state}}.

# 12. GOVERNING LAW AND JURISDICTION

# This Agreement shall be governed by and construed in accordance with the laws of the State of {{governing_state}}, without regard to its conflict of law provisions. The Parties consent to the exclusive jurisdiction of the courts located in {{governing_state}}.

# 13. GENERAL PROVISIONS

# 13.1 Entire Agreement
# This Agreement constitutes the entire agreement between the Parties and supersedes all prior agreements, whether written or oral.

# 13.2 Amendments
# This Agreement may only be amended in writing signed by both Parties. No oral modifications shall be binding.

# 13.3 Assignment
# Neither Party may assign this Agreement without the prior written consent of the other Party, except in connection with a merger, acquisition, or sale of substantially all assets.

# 13.4 Severability
# If any provision of this Agreement is found invalid or unenforceable, the remaining provisions shall continue in full force and effect.

# 13.5 Notices
# All notices under this Agreement shall be in writing and sent to the addresses listed above by certified mail, email, or other means that provide proof of delivery.

# 13.6 Waiver
# No waiver of any provision of this Agreement shall constitute a waiver of any other provision. Failure to enforce any provision shall not constitute a waiver.

# 13.7 Force Majeure
# Neither Party shall be liable for any failure or delay in performance due to circumstances beyond their reasonable control, including but not limited to acts of God, war, strikes, or natural disasters.

# IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

# SERVICE PROVIDER:

# _________________________________
# {{provider_name}}
# {{provider_title}}
# Date: {{signature_date}}


# CLIENT:

# _________________________________
# {{client_name}}
# {{client_title}}
# Date: {{signature_date}}
# """

#     def generate_contract(self, 
#                          contract_type: str, 
#                          user_data: Dict[str, str], 
#                          output_path: str,
#                          format_type: str = "docx") -> str:
#         """
#         Generate a contract from template with user data.
        
#         Args:
#             contract_type: Type of contract ("nda" or "service_agreement")
#             user_data: Dictionary with placeholder values
#             output_path: Path where the contract will be saved
#             format_type: Output format ("docx" or "txt")
            
#         Returns:
#             Path to the generated contract file
#         """
#         if contract_type not in self.templates:
#             raise ValueError(f"Unknown contract type: {contract_type}. Available: {list(self.templates.keys())}")
        
#         # Get template
#         template_content = self.templates[contract_type]
        
#         # Add default values for common fields if not provided
#         default_values = {
#             "signature_date": datetime.now().strftime("%B %d, %Y"),
#             "effective_date": datetime.now().strftime("%B %d, %Y"),
#         }
        
#         # Merge defaults with user data (user data takes precedence)
#         filled_data = {**default_values, **user_data}
        
#         # Fill in placeholders
#         contract_content = template_content
#         for key, value in filled_data.items():
#             placeholder = f"{{{{{key}}}}}"
#             contract_content = contract_content.replace(placeholder, str(value))
        
#         # Check for unfilled placeholders
#         unfilled = re.findall(r'\{\{([^}]+)\}\}', contract_content)
#         if unfilled:
#             print(f"Warning: The following placeholders were not filled: {', '.join(unfilled)}")
        
#         # Generate output file
#         if format_type == "docx":
#             return self._generate_docx(contract_content, output_path, contract_type)
#         else:
#             return self._generate_txt(contract_content, output_path)
    
#     def _generate_docx(self, content: str, output_path: str, contract_type: str) -> str:
#         """Generate a formatted DOCX file"""
#         doc = Document()
        
#         # Set document margins
#         sections = doc.sections
#         for section in sections:
#             section.top_margin = Inches(1)
#             section.bottom_margin = Inches(1)
#             section.left_margin = Inches(1)
#             section.right_margin = Inches(1)
        
#         # Add content with formatting
#         lines = content.strip().split('\n')
#         for line in lines:
#             line = line.strip()
#             if not line:
#                 doc.add_paragraph()
#                 continue
            
#             paragraph = doc.add_paragraph()
#             run = paragraph.add_run(line)
            
#             # Format titles (ALL CAPS lines)
#             if line.isupper() and len(line) > 3:
#                 run.bold = True
#                 run.font.size = Pt(14)
#                 paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             # Format section headers (lines starting with numbers)
#             elif line[0].isdigit() and '.' in line[:3]:
#                 run.bold = True
#                 run.font.size = Pt(12)
#             else:
#                 run.font.size = Pt(11)
        
#         # Ensure output directory exists
#         os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
#         # Save document
#         doc.save(output_path)
#         return output_path
    
#     def _generate_txt(self, content: str, output_path: str) -> str:
#         """Generate a plain text file"""
#         os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
#         with open(output_path, "w", encoding="utf-8") as f:
#             f.write(content)
        
#         return output_path
    
#     def get_required_fields(self, contract_type: str) -> list:
#         """Get list of required fields for a contract type"""
#         if contract_type not in self.templates:
#             raise ValueError(f"Unknown contract type: {contract_type}")
        
#         template = self.templates[contract_type]
#         placeholders = re.findall(r'\{\{([^}]+)\}\}', template)
        
#         # Remove duplicates and sort
#         return sorted(list(set(placeholders)))
    
#     def get_available_templates(self) -> Dict[str, str]:
#         """Get list of available contract templates with descriptions"""
#         return {
#             "nda": "Non-Disclosure Agreement (NDA) - Protects confidential information between parties",
#             "service_agreement": "Service Agreement - Defines terms for professional services between provider and client"
#         }


from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict
import os
import re

class ContractTemplateGenerator:
    """
    Rent agreement generator with comprehensive template.
    Supports filling placeholders with user data and generating downloadable contracts.
    """
    
    def __init__(self):
        self.templates = {
            "rent_agreement": self._get_rent_agreement_template(),
        }
    
    def _get_rent_agreement_template(self) -> str:
        """Rent Agreement Template"""
        return """
RENTAL AGREEMENT

This Rental Agreement ("Agreement") is made and entered into as of {{effective_date}}, by and between:

LANDLORD:
{{landlord_name}}
{{landlord_address}}
{{landlord_phone}}
{{landlord_email}}

and

TENANT:
{{tenant_name}}
{{tenant_address}}
{{tenant_phone}}
{{tenant_email}}

(collectively referred to as the "Parties")

WHEREAS, the Landlord is the lawful owner of the premises located at {{property_address}} ("Premises");

WHEREAS, the Tenant desires to lease the Premises from the Landlord for residential purposes;

NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, the Parties agree as follows:

1. PREMISES

The Landlord hereby leases to the Tenant, and the Tenant hereby leases from the Landlord, the Premises located at:

{{property_address}}

Property Type: {{property_type}}

The Premises shall be used exclusively for residential purposes by the Tenant and the following occupants only:
[List of additional occupants, if any]

2. TERM

2.1 Lease Term
The initial term of this Agreement shall commence on {{lease_start_date}} and end on {{lease_end_date}}.

2.2 Holdover Tenancy
If the Tenant remains in possession of the Premises after the expiration of the term with the Landlord's consent, the tenancy shall become month-to-month, subject to all other terms and conditions of this Agreement.

3. RENT

3.1 Monthly Rent
The Tenant shall pay to the Landlord as rent for the Premises the amount as agreed upon separately, payable in advance on the {{rent_due_date}} day of each calendar month.

3.2 Method of Payment
Rent shall be paid by the method agreed upon by both Parties.

3.3 Late Fees
If rent is not received by the Landlord by the due date, the Tenant shall be subject to late fees as specified in a separate agreement.

4. SECURITY DEPOSIT

4.1 Deposit Amount
Upon execution of this Agreement, the Tenant shall deposit with the Landlord a security deposit as agreed upon separately.

4.2 Use of Deposit
The security deposit may be used by the Landlord to remedy defaults in the payment of rent, to repair damages to the Premises caused by the Tenant, and to clean the Premises upon termination of tenancy.

4.3 Return of Deposit
The security deposit, less any deductions, shall be returned to the Tenant within 30 days after the termination of tenancy and delivery of possession to the Landlord.

5. UTILITIES

The following utilities are included in the rent:
{{utilities_included}}

All other utilities shall be the responsibility of the Tenant.

6. MAINTENANCE AND REPAIRS

6.1 Tenant Responsibilities
The Tenant shall:
   a) Keep the Premises clean and sanitary
   b) Dispose of all waste in a clean and safe manner
   c) Use all electrical, plumbing, and other facilities in a reasonable manner
   d) Notify the Landlord promptly of any needed repairs

6.2 Landlord Responsibilities
The Landlord shall maintain the Premises in a habitable condition and make all necessary repairs to comply with health and safety codes.

Maintenance Responsibility Details:
{{maintenance_responsibility}}

7. USE OF PREMISES

7.1 Residential Use Only
The Premises shall be used exclusively for residential purposes. No business or commercial activity shall be conducted on the Premises without the Landlord's prior written consent.

7.2 Compliance with Laws
The Tenant shall comply with all laws, rules, ordinances, and regulations affecting the Premises.

7.3 Quiet Enjoyment
The Tenant shall be entitled to quiet enjoyment of the Premises and shall not disturb other residents.

8. ALTERATIONS AND IMPROVEMENTS

The Tenant shall not make any alterations, additions, or improvements to the Premises without the Landlord's prior written consent. All alterations, additions, or improvements become part of the Premises and remain upon termination of tenancy.

9. INSURANCE

The Tenant is responsible for obtaining insurance for personal property. The Landlord's insurance does not cover the Tenant's personal property.

10. DEFAULT

10.1 Tenant Default
The following shall constitute default by the Tenant:
   a) Failure to pay rent when due
   b) Failure to perform any obligation under this Agreement
   c) Abandonment of the Premises
   d) Creating a nuisance or disturbing other residents

10.2 Landlord Remedies
Upon default by the Tenant, the Landlord may terminate this Agreement and pursue any remedies available under law.

11. TERMINATION

11.1 Termination Notice
Either party may terminate this Agreement by giving {{notice_period}} days written notice to the other party.

11.2 Surrender of Premises
Upon termination, the Tenant shall surrender the Premises in as good condition as when received, reasonable wear and tear excepted.

12. DISPUTE RESOLUTION

12.1 Good Faith Negotiation
In the event of any dispute arising out of or relating to this Agreement, the Parties shall first attempt to resolve the dispute through good faith negotiations for a period of thirty (30) days.

12.2 Mediation
If the dispute cannot be resolved through negotiation, the Parties agree to submit the dispute to mediation before a mutually agreed-upon mediator. The costs of mediation shall be shared equally by the Parties.

12.3 Arbitration
If mediation is unsuccessful, any remaining disputes shall be resolved through binding arbitration in accordance with the rules of the American Arbitration Association.

13. INSPECTIONS

The Landlord may enter the Premises at reasonable times to inspect, make repairs, or show the Premises to prospective tenants or purchasers, after giving 24 hours notice to the Tenant.

14. SUBLETTING AND ASSIGNMENT

The Tenant shall not sublet the Premises or assign this Agreement without the Landlord's prior written consent.

15. DAMAGES TO PREMISES

If the Premises are damaged by fire or other casualty, the Landlord shall repair the damage. If the Premises are rendered uninhabitable, rent shall abate until repairs are completed.

16. ENTIRE AGREEMENT

This Agreement constitutes the entire agreement between the Parties and supersedes all prior oral or written agreements relating to the Premises.

17. NOTICES

All notices under this Agreement shall be in writing and delivered to the addresses specified above, either in person or by certified mail.

18. SECURITY DEPOSIT ACKNOWLEDGMENT

The Tenant acknowledges receipt of the Landlord's security deposit disclosure statement, if required by state law.

19. LEAD-BASED PAINT DISCLOSURE

If the Premises were built before 1978, the Landlord has provided the Tenant with the federally approved lead-based paint disclosure information.

20. ADDITIONAL PROVISIONS

[Any additional provisions or special clauses agreed upon by the Parties]

IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

LANDLORD:

_________________________________
{{landlord_name}}
Date: {{signature_date}}

TENANT:

_________________________________
{{tenant_name}}
Date: {{signature_date}}

WITNESS:

_________________________________
Witness Name
Date: {{signature_date}}
"""

    def generate_contract(self, 
                         contract_type: str, 
                         user_data: Dict[str, str], 
                         output_path: str,
                         format_type: str = "docx") -> str:
        """
        Generate a rent agreement from template with user data.
        
        Args:
            contract_type: Type of contract ("rent_agreement")
            user_data: Dictionary with placeholder values
            output_path: Path where the contract will be saved
            format_type: Output format ("docx" or "txt")
            
        Returns:
            Path to the generated contract file
        """
        if contract_type not in self.templates:
            raise ValueError(f"Unknown contract type: {contract_type}. Available: {list(self.templates.keys())}")
        
        # Get template
        template_content = self.templates[contract_type]
        
        # Add default values for common fields if not provided
        default_values = {
            "signature_date": datetime.now().strftime("%B %d, %Y"),
            "effective_date": datetime.now().strftime("%B %d, %Y"),
        }
        
        # Merge defaults with user data (user data takes precedence)
        filled_data = {**default_values, **user_data}
        
        # Fill in placeholders
        contract_content = template_content
        for key, value in filled_data.items():
            placeholder = f"{{{{{key}}}}}"
            contract_content = contract_content.replace(placeholder, str(value))
        
        # Check for unfilled placeholders
        unfilled = re.findall(r'\{\{([^}]+)\}\}', contract_content)
        if unfilled:
            print(f"Warning: The following placeholders were not filled: {', '.join(unfilled)}")
        
        # Generate output file
        if format_type == "docx":
            return self._generate_docx(contract_content, output_path, contract_type)
        else:
            return self._generate_txt(contract_content, output_path)
    
    def _generate_docx(self, content: str, output_path: str, contract_type: str) -> str:
        """Generate a formatted DOCX file"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add content with formatting
        lines = content.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            
            # Format titles (ALL CAPS lines)
            if line.isupper() and len(line) > 3:
                run.bold = True
                run.font.size = Pt(14)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Format section headers (lines starting with numbers)
            elif line[0].isdigit() and '.' in line[:3]:
                run.bold = True
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(11)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _generate_txt(self, content: str, output_path: str) -> str:
        """Generate a plain text file"""
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return output_path
    
    def get_required_fields(self, contract_type: str) -> list:
        """Get list of required fields for a contract type"""
        if contract_type not in self.templates:
            raise ValueError(f"Unknown contract type: {contract_type}")
        
        template = self.templates[contract_type]
        placeholders = re.findall(r'\{\{([^}]+)\}\}', template)
        
        # Remove duplicates and sort
        return sorted(list(set(placeholders)))
    
    def get_available_templates(self) -> Dict[str, str]:
        """Get list of available contract templates with descriptions"""
        return {
            "rent_agreement": "Rent Agreement - Standard rental agreement between landlord and tenant"
        }

# Example usage
if __name__ == "__main__":
    generator = ContractTemplateGenerator()
    
    # Sample data for rent agreement
    sample_data = {
        "landlord_name": "John Smith",
        "landlord_address": "123 Owner Street, City, State 12345",
        "landlord_phone": "(555) 123-4567",
        "landlord_email": "john.smith@email.com",
        "tenant_name": "Jane Doe",
        "tenant_address": "456 Renter Avenue, City, State 12345",
        "tenant_phone": "(555) 987-6543",
        "tenant_email": "jane.doe@email.com",
        "property_address": "789 Rental Road, Apartment 4B, City, State 12345",
        "property_type": "Apartment",
        "lease_start_date": "January 1, 2024",
        "lease_end_date": "December 31, 2024",
        "rent_due_date": "1st",
        "utilities_included": "Water and trash collection. Tenant responsible for electricity, gas, and internet.",
        "maintenance_responsibility": "Landlord responsible for structural repairs and major systems. Tenant responsible for minor maintenance and keeping premises clean.",
        "notice_period": "30"
    }
    
    # Generate sample contract
    output_file = generator.generate_contract(
        contract_type="rent_agreement",
        user_data=sample_data,
        output_path="sample_rent_agreement.docx"
    )
    
    print(f"Generated contract: {output_file}")
    print(f"Required fields: {generator.get_required_fields('rent_agreement')}")