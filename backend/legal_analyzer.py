import re
import unicodedata
from PyPDF2 import PdfReader
from docx import Document
from dataclasses import dataclass, field 
from typing import Dict, List, Tuple, Any
import spacy
from spacy.matcher import Matcher
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from difflib import SequenceMatcher

 # Update the ComparisonResult dataclass to include new fields
@dataclass
class AnalysisResult:
    """Results from analyzing a single legal document"""
    clauses: Dict[str, List[Dict[str, Any]]] = field(default_factory=dict)
    entities: Dict[str, List[str]] = field(default_factory=dict)
    summary: str = ""
    statistics: Dict[str, int] = field(default_factory=dict)
    cleaned_text: str = ""
    signing_recommendation: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ComparisonResult:
    doc1_path: str
    doc2_path: str
    clause_differences: Dict[str, Dict[str, Any]]
    entity_differences: Dict[str, Dict[str, Any]]
    summary_comparison: Dict[str, Any]  # Changed from str to dict
    recommendation_comparison: Dict[str, Any]
    statistics_comparison: Dict[str, Any] = field(default_factory=dict)
    critical_differences: List[Dict[str, Any]] = field(default_factory=list)
    overall_similarity: Dict[str, Any] = field(default_factory=dict)

@dataclass
class AnalysisResult:
    clauses: Dict[str, List[Dict[str, Any]]] = field(default_factory=dict) # Use default_factory for mutable defaults
    entities: Dict[str, List[str]] = field(default_factory=dict)
    summary: str = ""
    statistics: Dict[str, int] = field(default_factory=dict)
    cleaned_text: str = ""
    signing_recommendation: Dict[str, Any] = field(default_factory=dict)

class LegalDocumentAnalyzer:
    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        try:
            self.nlp = spacy.load("en_core_web_lg")
            self.nlp.max_length = 10_000_000 # Increased max_length
        except OSError:
            raise ImportError(
                "Please install the spaCy English model: 'python -m spacy download en_core_web_lg' "
                "(recommended) or 'python -m spacy download en_core_web_sm')"
            )

        self.matcher = Matcher(self.nlp.vocab)
        self._add_custom_matcher_patterns()

        # Increased specificity for legal clauses
        self.legal_clauses = {
            "confidentiality": r"confidential(?:ity| information| agreement)",
            "indemnification": r"indemnif(?:y|ication|ies|ing)",
            "liability": r"liability|limitation of liability|disclaimer of liability",
            "termination": r"terminat(?:ion|e|es|ing)",
            "governing_law": r"governing law|choice of law|jurisdiction|venue",
            "amendment": r"amend(?:ment|ing|ed)",
            "dispute_resolution": r"dispute resolution|arbitration|mediation",
            "force_majeure": r"force majeure",
            "assignment": r"assign(?:ment|ability|s)",
            "warranty": r"warrant(?:y|ies|s)",
            "severability": r"severability"
        }

        # More detailed risk patterns with potentially higher weights
        self.risk_patterns = {
            "unlimited_liability": {
                "pattern": r"unlimited liability|all liability|full liability|no limitation of liability",
                "weight": -20, # Increased impact
                "description": "Unlimited or no limitation on liability"
            },
            "unilateral_changes": {
                "pattern": r"(?:may|can|will) (?:modify|change|alter|amend) (?:at any time|without notice|in its sole discretion|unilaterally)",
                "weight": -15, # Increased impact
                "description": "Party has unilateral modification rights"
            },
            "broad_indemnification": {
                "pattern": r"indemnif(?:y|ication) (?:against|from) any and all (?:claims|losses|damages)",
                "weight": -12, # A broad indemnification could be a risk depending on context
                "description": "Broad indemnification clause"
            },
            "automatic_renewal": {
                "pattern": r"automatically renew(?:s|ed|al)",
                "weight": -8,
                "description": "Automatic renewal without clear termination notice"
            },
            "no_assignment_without_consent": {
                "pattern": r"not assign(?:ed)? this agreement without.{1,50}prior (?:written )?consent",
                "weight": -7,
                "description": "Restriction on assignment without consent"
            },
            "broad_confidentiality": {
                "pattern": r"all(?: |\b)information(?: |\b)shall be considered confidential",
                "weight": -6,
                "description": "Overly broad confidentiality definition"
            },
            "non_negotiable": {
                "pattern": r"non-negotiable|not negotiable|as is|without recourse",
                "weight": -10,
                "description": "Non-negotiable terms"
            },
            "waiver_of_rights": {
                "pattern": r"waive(?:s|r of) (?:right|jury trial|class action)",
                "weight": -15,
                "description": "Waiver of significant rights (e.g., jury trial)"
            },
            "perpetual_obligations": {
                "pattern": r"(?:perpetual|eternal|indefinite|survive termination)",
                "weight": -8,
                "description": "Perpetual obligations after termination"
            }
        }

        # More detailed favorable patterns with potentially higher weights
        self.favorable_patterns = {
            "mutual_termination": {
                "pattern": r"(?:either|both|any) part(?:y|ies) may terminate",
                "weight": 12, # Increased impact
                "description": "Mutual termination rights for all parties"
            },
            "limited_liability": {
                "pattern": r"liability (?:limited to|shall not exceed) (?:[$€£]\d{1,3}(?:,\d{3})*(?:\.\d{2})?|the amount of this agreement)",
                "weight": 15, # Strong positive impact
                "description": "Liability is clearly limited to a specific amount"
            },
            "reasonable_notice_period": {
                "pattern": r"(?:notice period|notice of \d+ (?:business )?(?:day|week|month|year)s) prior to termination",
                "weight": 8,
                "description": "Clearly defined reasonable notice period for termination"
            },
            "clear_dispute_resolution": {
                "pattern": r"dispute resolution clause|(?:binding )?arbitration|(?:mandatory )?mediation",
                "weight": 10,
                "description": "Clear and defined alternative dispute resolution mechanism"
            },
            "mutual_confidentiality": {
                "pattern": r"(?:both|all|either|respective) part(?:y|ies).{1,50}confidential(?:ity)? obligations",
                "weight": 7,
                "description": "Mutual confidentiality obligations"
            },
            "governing_law_defined": {
                "pattern": r"governed by and construed in accordance with the laws of (?:the State of )?[A-Za-z ]+",
                "weight": 8,
                "description": "Clearly defined governing law and jurisdiction"
            },
            "right_to_cure": {
                "pattern": r"(?:right to cure|opportunity to cure) (?:breach)?",
                "weight": 7,
                "description": "Opportunity to cure breaches before termination"
            }
        }

    def print_debug(self, message: str):
        if self.verbose:
            print(f"[DEBUG] {message}")

    def _add_custom_matcher_patterns(self):
        # Define patterns for common legal entities that spaCy might miss or misclassify.
        # These patterns are designed to be more specific to legal contexts.

        # --- Document Identifiers ---
        self.matcher.add("LEGAL_DOC_TYPE", [[{"LOWER": "this"}, {"LOWER": {"IN": ["agreement", "contract", "memorandum", "indenture", "deed"]}}]])
        self.matcher.add("CONTRACT_TITLE", [[{"POS": "PROPN", "OP": "*"}, {"LOWER": "agreement"}],
                                             [{"POS": "PROPN", "OP": "*"}, {"LOWER": "contract"}],
                                             [{"IS_TITLE": True, "OP": "+"}, {"LOWER": {"IN": ["agreement", "contract"]}}]]) # e.g. "Software License Agreement"

        # --- Parties to the Agreement ---
        self.matcher.add("LEGAL_PARTY_ROLE", [
            [{"LOWER": {"IN": ["client", "contractor", "vendor", "licensor", "licensee", "supplier", "grantor", "grantee", "employer", "employee", "customer"]}}],
            [{"LOWER": "party"}, {"TEXT": {"REGEX": "[A-Z]"}}], # e.g., "Party A", "Party B"
            [{"LOWER": "parties"}],
            [{"LOWER": "the"}, {"LOWER": {"IN": ["undersigned", "company", "corporation", "inc.", "llc", "ltd.", "llp", "pte ltd", "plc", "gmbh", "ag", "sa", "co.", "inc"]}}] # Added more common suffixes
        ])
        # Add a pattern for specific company names that aren't picked up by ORG
        self.matcher.add("COMPANY_NAME_PAT", [[{"POS": "PROPN", "OP": "+"}, {"LOWER": {"IN": ["inc.", "llc", "corp.", "ltd.", "co."]}}]])


        # --- Dates and Durations ---
        self.matcher.add("EFFECTIVE_DATE", [[{"LOWER": "effective"}, {"LOWER": "as"}, {"LOWER": "of"}, {"ENT_TYPE": "DATE"}]])
        self.matcher.add("AGREEMENT_DATE", [[{"LOWER": "date"}, {"LOWER": "of"}, {"LOWER": "this"}, {"LOWER": "agreement"}, {"IS_PUNCT": False, "OP": "*"}, {"ENT_TYPE": "DATE"}]])
        self.matcher.add("LEGAL_DURATION", [[{"ENT_TYPE": "CARDINAL"}, {"LOWER": {"IN": ["day", "days", "week", "weeks", "month", "months", "year", "years", "business days", "calendar days"]}}]])

        # --- Monetary Terms ---
        self.matcher.add("LEGAL_MONEY", [
            [{"TEXT": {"REGEX": r"\$|€|£|¥"}}, {"IS_DIGIT": True, "OP": "+"}, {"TEXT": ".", "OP": "?"}, {"IS_DIGIT": True, "OP": "*"}], # Currency symbol then digits
            [{"TEXT": {"REGEX": r"\d{1,3}(?:,\d{3})*(?:\.\d{2})?"}}, {"LOWER": {"IN": ["usd", "eur", "gbp", "dollars", "euros", "pounds", "yen"]}, "OP": "+"}], # Digits then currency code/word
            [{"ENT_TYPE": "CARDINAL"}, {"LOWER": {"IN": ["dollars", "euros", "pounds"]}}] # "One thousand dollars"
        ])

        # --- Governing Law and Jurisdiction ---
        self.matcher.add("GOVERNING_LAW_LOC", [[{"LOWER": {"IN": ["laws", "jurisdiction"]}}, {"LOWER": "of"}, {"ENT_TYPE": "GPE"}]])
        self.matcher.add("STATE_NAME", [[{"LOWER": "state"}, {"LOWER": "of"}, {"POS": "PROPN", "OP": "+"}]])
        self.matcher.add("COMMONWEALTH_NAME", [[{"LOWER": "commonwealth"}, {"LOWER": "of"}, {"POS": "PROPN", "OP": "+"}]])
        # Broader pattern for governing law phrases
        self.matcher.add("GOVERNING_LAW_PHRASE", [[{"LOWER": "governed"}, {"LOWER": "by"}, {"LOWER": "and"}, {"LOWER": "construed"}, {"LOWER": "in"}, {"LOWER": "accordance"}, {"LOWER": "with"}, {"LOWER": "the"}, {"LOWER": "laws"}, {"LOWER": "of"}, {"ENT_TYPE": "GPE", "OP": "*"}, {"POS": "PROPN", "OP": "+"}, {"LOWER": "state", "OP": "?"}]])


        # --- Legal Processes/Concepts (often misclassified by general NER) ---
        self.matcher.add("LEGAL_CONCEPT", [[{"LOWER": {"IN": ["arbitration", "mediation", "litigation", "injunction", "subrogation", "negotiation", "termination", "indemnification", "confidentiality", "governing law", "jurisdiction", "amendment", "waiver", "notice", "breach", "remedies", "damages", "force majeure", "severability", "assignment", "warranty", "exclusive", "non-exclusive", "royalty", "licence"]}}]])
        self.matcher.add("LEGAL_CONCEPT", [[{"LOWER": "limitation"}, {"LOWER": "of"}, {"LOWER": "liability"}]])
        self.matcher.add("LEGAL_CONCEPT", [[{"LOWER": "dispute"}, {"LOWER": "resolution"}]])
        self.matcher.add("LEGAL_CONCEPT", [[{"LOWER": "intellectual"}, {"LOWER": "property"}]])


    def clean_text(self, text: str) -> str:
        replacements = {
            'Ɵ': 't', 'Ʃ': 's', 'ƚ': 'l', 'ƭ': 't',
            'ﬁ': 'fi', 'ﬂ': 'fl', 'ﬀ': 'ff', 'ﬃ': 'ffi',
            'conﬁdenƟal': 'confidential',
            'Terminaton': 'Termination',
            'Compensa ton': 'Compensation',
            'informa ton': 'information',
            'arbitra ton': 'arbitration',
            'Arbitra': 'Arbitration'
        }
        for bad_char, good_char in replacements.items():
            text = text.replace(bad_char, good_char)
        text = unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode('utf-8')

        # Replace multiple spaces with a single space, but preserve newlines.
        text = re.sub(r'[ \t]+', ' ', text) # Replace multiple spaces/tabs with single space
        text = re.sub(r'\n{2,}', '\n\n', text).strip() # Reduce multiple newlines to max two for paragraphs

        return text

    def load_document(self, file_path: str) -> str:
        if file_path.endswith('.pdf'):
            return self._load_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self._load_docx(file_path)
        else:
            return self._load_text(file_path)

    def _load_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return self.clean_text(text)

    def _load_docx(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            # When extracting from DOCX, ensure paragraphs are separated by newlines
            return self.clean_text("\n".join([para.text for para in doc.paragraphs]))
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")

    def _load_text(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return self.clean_text(file.read())
        except Exception as e:
            raise ValueError(f"Error reading text file: {str(e)}")

    def identify_clauses(self, text: str) -> Dict[str, List[Dict[str, Any]]]:
        found_clauses = {}
        
        for name, pattern in self.legal_clauses.items():
            matches = re.finditer(pattern, text, re.IGNORECASE)
            clause_matches = []
            for match in matches:
                # Calculate line number
                line_number = text.count('\n', 0, match.start()) + 1
                clause_matches.append({
                    "text": match.group(),
                    "positions": [match.start(), match.end()],
                    "line_number": line_number
                })
            # Sort by line number for consistent output
            found_clauses[name] = sorted(clause_matches, key=lambda x: x['line_number'])
        return found_clauses


    def extract_entities(self, text: str) -> Dict[str, List[str]]:
        doc = self.nlp(text)
        temp_entities = {}

        for ent in doc.ents:
            temp_entities.setdefault(ent.label_, []).append(ent.text)

        matches = self.matcher(doc)
        for match_id, start, end in matches:
            span = doc[start:end]
            label = self.nlp.vocab.strings[match_id]
            temp_entities.setdefault(label, []).append(span.text)

        final_entities: Dict[str, List[str]] = {
            "CONTRACT_PARTIES": [],
            "CONTRACT_DATES": [],
            "MONEY_AMOUNTS": [],
            "LOCATIONS": [],
            "ORGANIZATIONS": [],
            "PERSONS": [],
            "LEGAL_CONCEPTS": [],
            "DOCUMENT_TYPES": [],
            "DURATIONS": [],
            "GOVERNING_LAW_LOCATIONS": [],
            "CONTRACT_TITLES": [],
            "CARDINAL_NUMBERS": [],
            "ORDINAL_NUMBERS": [],
            "COMPANY_NAMES": [] # New category for explicit company names
        }

        # Priority 1: Custom Matcher entities
        for label_key in [
            "LEGAL_DOC_TYPE", "CONTRACT_TITLE", "LEGAL_PARTY_ROLE", "COMPANY_NAME_PAT",
            "EFFECTIVE_DATE", "AGREEMENT_DATE", "LEGAL_DURATION",
            "LEGAL_MONEY", "GOVERNING_LAW_LOC", "STATE_NAME", "COMMONWEALTH_NAME", "GOVERNING_LAW_PHRASE",
            "LEGAL_CONCEPT"
        ]:
            if label_key in temp_entities:
                if label_key == "LEGAL_DOC_TYPE":
                    final_entities["DOCUMENT_TYPES"].extend(temp_entities[label_key])
                elif label_key == "CONTRACT_TITLE":
                    final_entities["CONTRACT_TITLES"].extend(temp_entities[label_key])
                elif label_key == "LEGAL_PARTY_ROLE":
                    final_entities["CONTRACT_PARTIES"].extend(temp_entities[label_key])
                elif label_key == "COMPANY_NAME_PAT":
                    final_entities["COMPANY_NAMES"].extend(temp_entities[label_key])
                elif label_key in ["EFFECTIVE_DATE", "AGREEMENT_DATE"]:
                    final_entities["CONTRACT_DATES"].extend(temp_entities[label_key])
                elif label_key == "LEGAL_DURATION":
                    final_entities["DURATIONS"].extend(temp_entities[label_key])
                elif label_key == "LEGAL_MONEY":
                    final_entities["MONEY_AMOUNTS"].extend(temp_entities[label_key])
                elif label_key in ["GOVERNING_LAW_LOC", "STATE_NAME", "COMMONWEALTH_NAME", "GOVERNING_LAW_PHRASE"]:
                    final_entities["GOVERNING_LAW_LOCATIONS"].extend(temp_entities[label_key])
                elif label_key == "LEGAL_CONCEPT":
                    final_entities["LEGAL_CONCEPTS"].extend(temp_entities[label_key])

        # Priority 2: General spaCy NER entities, with filtering and reclassification
        # ORGANIZATIONS (ORG)
        if "ORG" in temp_entities:
            for org_candidate in temp_entities["ORG"]:
                lowered_org = org_candidate.lower()
                # Skip if already captured by more specific COMPANY_NAMES or LEGAL_PARTY_ROLE
                if any(org_candidate in c for c in final_entities["COMPANY_NAMES"]) or \
                   any(org_candidate in p for p in final_entities["CONTRACT_PARTIES"]):
                    continue
                # Exclude if it's a known legal role or concept or document title
                if any(keyword in lowered_org for keyword in
                       [lc.lower() for lc in final_entities["LEGAL_CONCEPTS"]] +
                       [lp.lower() for lp in final_entities["CONTRACT_PARTIES"]] +
                       [lt.lower() for lt in final_entities["CONTRACT_TITLES"]] +
                       [ldt.lower() for ldt in final_entities["DOCUMENT_TYPES"]]):
                    self.print_debug(f"Skipping ORG '{org_candidate}' (likely a legal concept/role/title).")
                    continue
                # Specific common misclassifications
                if lowered_org in ["client", "contractor", "developer", "vendor", "agreement", "resolution"]:
                    self.print_debug(f"Reclassifying ORG '{org_candidate}' to CONTRACT_PARTIES/LEGAL_CONCEPTS.")
                    if lowered_org in ["client", "contractor", "developer", "vendor"]:
                        final_entities["CONTRACT_PARTIES"].append(org_candidate)
                    elif lowered_org in ["agreement", "resolution"]:
                        final_entities["LEGAL_CONCEPTS"].append(org_candidate)
                    continue
                # If it contains a zip code, it might be an address, not an organization.
                if re.search(r'\b\d{5}(?:-\d{4})?\b', org_candidate) and (len(org_candidate.split()) < 4 or any(c in org_candidate.lower() for c in ['street', 'road', 'avenue'])):
                    self.print_debug(f"Reclassifying ORG '{org_candidate}' to LOCATION (due to ZIP code/address hint).")
                    final_entities["LOCATIONS"].append(org_candidate)
                    continue
                # Otherwise, keep as ORG if it looks like a valid organization name
                if re.search(r'\b(?:inc|corp|llc|ltd|gmbh|ag|sa|co\.)\b', lowered_org) or \
                   (org_candidate[0].isupper() and ' ' in org_candidate and not any(digit.isdigit() for digit in org_candidate) and len(org_candidate.split()) > 1):
                    final_entities["ORGANIZATIONS"].append(org_candidate)
                else:
                    self.print_debug(f"Skipping ORG '{org_candidate}' (does not resemble typical organization name).")

        # PERSONS
        if "PERSON" in temp_entities:
            for person_candidate in temp_entities["PERSON"]:
                lowered_person = person_candidate.lower()
                if any(keyword in lowered_person for keyword in
                       [lp.lower() for lp in final_entities["CONTRACT_PARTIES"]] +
                       [lc.lower() for lc in final_entities["LEGAL_CONCEPTS"]]):
                    self.print_debug(f"Skipping PERSON '{person_candidate}' (likely a legal role/concept).")
                    continue
                if lowered_person in ["termination", "information", "compensation", "arbitration", "developer"]:
                    self.print_debug(f"Reclassifying PERSON '{person_candidate}' to LEGAL_CONCEPTS.")
                    final_entities["LEGAL_CONCEPTS"].append(person_candidate)
                    continue
                if len(person_candidate.split()) > 1 or re.match(r'^[A-Z]\.?\s?[A-Z][a-z]+$', person_candidate):
                    final_entities["PERSONS"].append(person_candidate)
                else:
                    self.print_debug(f"Skipping PERSON '{person_candidate}' (does not resemble typical person name).")

        # LOCATIONS (GPE)
        if "GPE" in temp_entities:
            for gpe_candidate in temp_entities["GPE"]:
                lowered_gpe = gpe_candidate.lower()
                if any(gpe_candidate in loc for loc in final_entities["GOVERNING_LAW_LOCATIONS"]):
                    continue
                if any(st in lowered_gpe for st in ["drive", "road", "street", "avenue", "boulevard"]):
                    self.print_debug(f"Skipping GPE '{gpe_candidate}' (likely just a street type).")
                    continue
                if lowered_gpe == "client":
                    final_entities["CONTRACT_PARTIES"].append(gpe_candidate)
                    continue
                final_entities["LOCATIONS"].append(gpe_candidate)

        # DATES
        if "DATE" in temp_entities:
            for date_candidate in temp_entities["DATE"]:
                if not re.search(r'\b\d{5}(?:-\d{4})?\b', date_candidate) and \
                   not any(date_candidate in cd for cd in final_entities["CONTRACT_DATES"]):
                    final_entities["CONTRACT_DATES"].append(date_candidate)

        # MONEY
        if "MONEY" in temp_entities:
            for money_candidate in temp_entities["MONEY"]:
                if not any(money_candidate in lm for lm in final_entities["MONEY_AMOUNTS"]):
                    final_entities["MONEY_AMOUNTS"].append(money_candidate)

        # CARDINAL and ORDINAL Numbers:
        if "CARDINAL" in temp_entities:
            for cardinal_candidate in temp_entities["CARDINAL"]:
                is_part_of_other_entity = False
                for entity_list in [final_entities["DURATIONS"], final_entities["MONEY_AMOUNTS"]]:
                    if any(cardinal_candidate in item for item in entity_list):
                        is_part_of_other_entity = True
                        break
                if not is_part_of_other_entity:
                    final_entities["CARDINAL_NUMBERS"].append(cardinal_candidate)

        if "ORDINAL" in temp_entities:
            final_entities["ORDINAL_NUMBERS"].extend(temp_entities["ORDINAL"])

        # Discard inherently irrelevant general spaCy entities for legal docs
        for discard_type in ["WORK_OF_ART", "PRODUCT", "EVENT", "FAC", "LANGUAGE"]:
            if discard_type in temp_entities:
                self.print_debug(f"Discarding {discard_type}: {temp_entities[discard_type]}")

        if "LAW" in temp_entities:
            for law_text in temp_entities["LAW"]:
                if "act" in law_text.lower() or "code" in law_text.lower():
                    final_entities.setdefault("LEGAL_STATUTES", []).append(law_text)
                else:
                    final_entities["LEGAL_CONCEPTS"].append(law_text)
        if "NORP" in temp_entities:
            final_entities.setdefault("NATIONALITIES_GROUPS", []).extend(temp_entities["NORP"])

        # --- Step 4: Deduplicate and Final Clean Up ---
        final_cleaned_entities = {}
        for key, value_list in final_entities.items():
            cleaned_list = list(set(value_list))
            if cleaned_list:
                final_cleaned_entities[key] = cleaned_list

        return final_cleaned_entities


    def summarize_text(self, text: str, sentences_count: int = 3) -> str:
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, sentences_count)
        return " ".join(str(sentence) for sentence in summary)

    def calculate_signing_recommendation(self, text: str, clause_data: Dict[str, Any], extracted_entities: Dict[str, List[str]]) -> Dict[str, Any]:
        # Start with a neutral base score
        # base_score = 50
        base_score = 60
        score = base_score
        findings = {
            "favorable_factors": [],
            "risk_factors": []
        }

        # Factor 1: Risk Patterns
        for name, risk_info in self.risk_patterns.items():
            matches = list(re.finditer(risk_info["pattern"], text, re.IGNORECASE))
            if matches:
                # Apply weight based on number of matches, but cap impact
                impact = risk_info["weight"] * min(len(matches), 3) # Cap impact for multiple occurrences
                score += impact
                findings["risk_factors"].append({
                    "type": name,
                    "description": risk_info["description"],
                    "weight": impact,
                    "matches": len(matches),
                    "examples": [text[max(0, m.start()-50):min(len(text), m.end()+50)] for m in matches[:2]]
                })

        # Factor 2: Favorable Patterns
        for name, favorable_info in self.favorable_patterns.items():
            matches = list(re.finditer(favorable_info["pattern"], text, re.IGNORECASE))
            if matches:
                # Apply weight based on number of matches, but cap impact
                impact = favorable_info["weight"] * min(len(matches), 3) # Cap impact for multiple occurrences
                score += impact
                findings["favorable_factors"].append({
                    "type": name,
                    "description": favorable_info["description"],
                    "weight": impact,
                    "matches": len(matches),
                    "examples": [text[max(0, m.start()-50):min(len(text), m.end()+50)] for m in matches[:2]]
                })

        # Factor 3: Presence/Absence of Key Clauses (from `identify_clauses`)
        critical_clauses = {
            "indemnification": "Missing indemnification clause",
            "governing_law": "Missing governing law clause",
            "termination": "Missing termination clause",
            "dispute_resolution": "Missing dispute resolution clause"
        }
        missing_clauses = []
        for clause_name, description in critical_clauses.items():
            # Check if list is empty for the clause name
            if not clause_data.get(clause_name):
                missing_clauses.append(clause_name)
                score -= 10 # Stronger penalty for missing critical clauses
                findings["risk_factors"].append({
                    "type": f"missing_{clause_name}",
                    "description": description,
                    "weight": -10
                })
            else:
                # If a critical clause IS present, add a slight positive
                score += 3
                findings["favorable_factors"].append({
                    "type": f"present_{clause_name}",
                    "description": f"Presence of {clause_name} clause",
                    "weight": 3
                })


        # Factor 4: Document Complexity / Readability
        word_count = len(text.split())
        if word_count < 200: # Very short document might indicate missing info
            score -= 10
            findings["risk_factors"].append({"type": "short_document", "description": "Document is very short, potentially incomplete", "weight": -10})
        elif word_count > 5000: # Very long documents can be complex
            score -= 5
            findings["risk_factors"].append({"type": "long_document", "description": "Document is very long, potentially complex", "weight": -5})

        avg_word_length = sum(len(word) for word in text.split()) / max(1, word_count)
        if avg_word_length > 7: # High average word length suggests complex language
            score -= (avg_word_length - 7) * 2 # More dynamic penalty
            findings["risk_factors"].append({
                "type": "complex_language",
                "description": f"Document uses complex language (avg word length: {avg_word_length:.2f})",
                "weight": -(round((avg_word_length - 7) * 2))
            })
        else: # Reward for simpler language
            score += 2


        # Factor 5: Entity Density (Example: presence of defined parties, money)
        if extracted_entities.get("CONTRACT_PARTIES"):
            score += 5 # Good to have clear parties
            findings["favorable_factors"].append({"type": "clear_parties", "description": "Clear identification of contract parties", "weight": 5})
        else:
            score -= 8
            findings["risk_factors"].append({"type": "unclear_parties", "description": "Ambiguous or missing contract parties", "weight": -8})

        if extracted_entities.get("MONEY_AMOUNTS"):
            score += 3 # Good to have money amounts defined
            findings["favorable_factors"].append({"type": "defined_money", "description": "Presence of monetary terms", "weight": 3})
        else:
            # Only a risk if it's a contract where money is expected
            # For a general analyzer, this is less critical unless specified
            pass

        if extracted_entities.get("GOVERNING_LAW_LOCATIONS"):
            score += 4 # Good to have governing law
            findings["favorable_factors"].append({"type": "defined_governing_law", "description": "Clearly defined governing law", "weight": 4})
        else:
            score -= 7
            findings["risk_factors"].append({"type": "missing_governing_law_entity", "description": "Governing law entity not explicitly found", "weight": -7})


        # Cap the score to stay within 0-100
        final_score = max(0, min(100, score))

        if final_score >= 80:
            recommendation = "Highly Favorable - Confident to sign after normal review."
        elif final_score >= 65: # Adjusted thresholds
            recommendation = "Favorable - Consider signing with normal review."
        elif final_score >= 50: # Adjusted thresholds
            recommendation = "Moderately Favorable - Review carefully; minor issues may exist."
        elif final_score >= 35: # Adjusted thresholds
            recommendation = "Neutral to Potentially Unfavorable - Consult legal professional for review."
        elif final_score >= 20:
            recommendation = "Potentially Unfavorable - Negotiate changes; significant risks identified."
        else:
            recommendation = "Highly Unfavorable - Do not sign; significant revisions or rejection recommended."

        return {
            "percentage": final_score,
            "recommendation": recommendation,
            "findings": findings,
            "missing_clauses": missing_clauses
        }

    def analyze(self, file_path: str) -> AnalysisResult:
        raw_text = self.load_document(file_path)
        cleaned_text = raw_text # Keep raw_text for line numbering based on original structure
        clauses = self.identify_clauses(cleaned_text)
        entities = self.extract_entities(cleaned_text) # Extract entities first
        return AnalysisResult(
            clauses=clauses,
            entities=entities, # Pass extracted entities
            summary=self.summarize_text(cleaned_text),
            statistics={
                "word_count": len(cleaned_text.split()),
                "char_count": len(cleaned_text),
                "paragraph_count": len(re.findall(r'\n\s*\n', raw_text)) # Use raw_text for paragraph count
            },
            cleaned_text=cleaned_text,
            # Pass extracted_entities to calculate_signing_recommendation
            signing_recommendation=self.calculate_signing_recommendation(cleaned_text, clauses, entities)
        )
   

    
    from typing import List, Set

    def compare_documents(self, file_path1: str, file_path2: str) -> ComparisonResult:
        """
        Enhanced document comparison with detailed clause and entity analysis.
        Provides actionable insights on differences between two legal documents.
        """
        # Analyze both documents
        result1 = self.analyze(file_path1)
        result2 = self.analyze(file_path2)

        # ===== Enhanced Clause Comparison =====
        clause_diff = {}
        all_clause_types = set(result1.clauses.keys()).union(result2.clauses.keys())
        
        for clause_name in all_clause_types:
            doc1_clauses = result1.clauses.get(clause_name, [])
            doc2_clauses = result2.clauses.get(clause_name, [])
            
            # Extract just the text for comparison
            doc1_texts = [c["text"] for c in doc1_clauses]
            doc2_texts = [c["text"] for c in doc2_clauses]
            
            # Determine status
            if doc1_clauses and doc2_clauses:
                status = "present_in_both"
                # Calculate similarity between clause texts
                similarity = self._calculate_text_similarity(
                    " ".join(doc1_texts), 
                    " ".join(doc2_texts)
                )
            elif doc1_clauses and not doc2_clauses:
                status = "only_in_doc1"
                similarity = 0.0
            elif doc2_clauses and not doc1_clauses:
                status = "only_in_doc2"
                similarity = 0.0
            else:
                continue  # Skip if absent in both
            
            clause_diff[clause_name] = {
                "status": status,
                "similarity": round(similarity * 100, 2),
                "doc1": {
                    "count": len(doc1_clauses),
                    "texts": doc1_texts,
                    "line_numbers": [c["line_number"] for c in doc1_clauses]
                },
                "doc2": {
                    "count": len(doc2_clauses),
                    "texts": doc2_texts,
                    "line_numbers": [c["line_number"] for c in doc2_clauses]
                },
                "analysis": self._analyze_clause_difference(
                    clause_name, doc1_texts, doc2_texts
                )
            }

        # ===== Enhanced Entity Comparison =====
        entity_diff = {}
        all_entity_types = set(result1.entities.keys()).union(result2.entities.keys())
        
        for entity_name in all_entity_types:
            doc1_entities = set(result1.entities.get(entity_name, []))
            doc2_entities = set(result2.entities.get(entity_name, []))
            
            common = doc1_entities.intersection(doc2_entities)
            only_in_1 = doc1_entities - doc2_entities
            only_in_2 = doc2_entities - doc1_entities
            
            # Calculate Jaccard similarity
            union = doc1_entities.union(doc2_entities)
            similarity = len(common) / len(union) if union else 1.0
            
            entity_diff[entity_name] = {
                "similarity": round(similarity * 100, 2),
                "common": sorted(list(common)),
                "only_in_doc1": sorted(list(only_in_1)),
                "only_in_doc2": sorted(list(only_in_2)),
                "doc1_count": len(doc1_entities),
                "doc2_count": len(doc2_entities),
                "significance": self._assess_entity_significance(entity_name, only_in_1, only_in_2)
            }

        # ===== Enhanced Summary Comparison =====
        summary_similarity = self._calculate_text_similarity(
            result1.summary, 
            result2.summary
        )
        
        summary_comp = {
            "doc1_summary": result1.summary,
            "doc2_summary": result2.summary,
            "similarity": round(summary_similarity * 100, 2),
            "key_differences": self._extract_summary_differences(
                result1.summary, 
                result2.summary
            )
        }

        # ===== Risk Score Comparison =====
        rec_comp = {
            "doc1": {
                "score": result1.signing_recommendation["percentage"],
                "recommendation": result1.signing_recommendation["recommendation"],
                "risk_factors": result1.signing_recommendation["findings"]["risk_factors"],
                "favorable_factors": result1.signing_recommendation["findings"]["favorable_factors"]
            },
            "doc2": {
                "score": result2.signing_recommendation["percentage"],
                "recommendation": result2.signing_recommendation["recommendation"],
                "risk_factors": result2.signing_recommendation["findings"]["risk_factors"],
                "favorable_factors": result2.signing_recommendation["findings"]["favorable_factors"]
            },
            "score_difference": result2.signing_recommendation["percentage"] - 
                            result1.signing_recommendation["percentage"],
            "which_is_better": self._determine_better_document(
                result1.signing_recommendation, 
                result2.signing_recommendation
            )
        }

        # ===== Statistical Comparison =====
        stats_comp = {
            "doc1": result1.statistics,
            "doc2": result2.statistics,
            "differences": {
                "word_count_diff": result2.statistics["word_count"] - result1.statistics["word_count"],
                "char_count_diff": result2.statistics["char_count"] - result1.statistics["char_count"],
                "paragraph_count_diff": result2.statistics["paragraph_count"] - result1.statistics["paragraph_count"]
            }
        }

        # ===== Critical Differences Analysis =====
        critical_diffs = self._identify_critical_differences(
            clause_diff, 
            entity_diff, 
            rec_comp
        )

        # Return enhanced comparison result
        return ComparisonResult(
            doc1_path=file_path1,
            doc2_path=file_path2,
            clause_differences=clause_diff,
            entity_differences=entity_diff,
            summary_comparison=summary_comp,
            recommendation_comparison=rec_comp,
            statistics_comparison=stats_comp,
            critical_differences=critical_diffs,
            overall_similarity=self._calculate_overall_similarity(
                clause_diff, 
                entity_diff, 
                summary_similarity
            )
        )


    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Calculate similarity ratio between two texts using SequenceMatcher."""
        if not text1 and not text2:
            return 1.0
        if not text1 or not text2:
            return 0.0
        return SequenceMatcher(None, text1.lower(), text2.lower()).ratio()


    def _analyze_clause_difference(self, clause_name: str, doc1_texts: List[str], 
                                doc2_texts: List[str]) -> str:
        """Provide human-readable analysis of clause differences."""
        if not doc1_texts and not doc2_texts:
            return "Clause absent in both documents"
        elif not doc1_texts:
            return f"Clause only present in Document 2 ({len(doc2_texts)} occurrence(s))"
        elif not doc2_texts:
            return f"Clause only present in Document 1 ({len(doc1_texts)} occurrence(s))"
        else:
            count_diff = len(doc2_texts) - len(doc1_texts)
            if count_diff > 0:
                return f"Clause appears more frequently in Doc 2 (+{count_diff} occurrence(s))"
            elif count_diff < 0:
                return f"Clause appears more frequently in Doc 1 ({count_diff} occurrence(s))"
            else:
                # Same count, check text similarity
                similarity = self._calculate_text_similarity(
                    " ".join(doc1_texts), 
                    " ".join(doc2_texts)
                )
                if similarity > 0.9:
                    return "Clauses are substantially similar"
                elif similarity > 0.7:
                    return "Clauses have moderate differences in wording"
                else:
                    return "Clauses have significant wording differences"


    def _assess_entity_significance(self, entity_type: str, only_in_1: Set[str], 
                                    only_in_2: Set[str]) -> str:
        """Assess the significance of entity differences."""
        critical_types = {
            "CONTRACT_PARTIES", "MONEY_AMOUNTS", "GOVERNING_LAW_LOCATIONS",
            "CONTRACT_DATES", "COMPANY_NAMES"
        }
        
        if entity_type in critical_types:
            if only_in_1 or only_in_2:
                return "HIGH - Critical entity differences detected"
            else:
                return "LOW - Critical entities match"
        else:
            if len(only_in_1) + len(only_in_2) > 5:
                return "MEDIUM - Multiple entity differences"
            else:
                return "LOW - Minor entity differences"


    def _extract_summary_differences(self, summary1: str, summary2: str) -> List[str]:
        """Extract key differences between summaries."""
        differences = []
        
        # Split into sentences
        sentences1 = set(sent.strip() for sent in summary1.split('.') if sent.strip())
        sentences2 = set(sent.strip() for sent in summary2.split('.') if sent.strip())
        
        unique_to_1 = sentences1 - sentences2
        unique_to_2 = sentences2 - sentences1
        
        if unique_to_1:
            differences.append(f"Doc 1 unique points: {'; '.join(list(unique_to_1)[:2])}")
        if unique_to_2:
            differences.append(f"Doc 2 unique points: {'; '.join(list(unique_to_2)[:2])}")
        
        return differences if differences else ["Summaries are very similar"]


    def _determine_better_document(self, rec1: dict, rec2: dict) -> dict:
        """Determine which document is more favorable based on multiple factors."""
        score1 = rec1["percentage"]
        score2 = rec2["percentage"]
        
        risk_count1 = len(rec1["findings"]["risk_factors"])
        risk_count2 = len(rec2["findings"]["risk_factors"])
        
        if score2 > score1 + 10:  # Significant difference
            return {
                "better_document": "Document 2",
                "reason": f"Significantly higher score ({score2} vs {score1})",
                "recommendation": "Consider using Document 2 as the base"
            }
        elif score1 > score2 + 10:
            return {
                "better_document": "Document 1",
                "reason": f"Significantly higher score ({score1} vs {score2})",
                "recommendation": "Consider using Document 1 as the base"
            }
        else:
            if risk_count1 < risk_count2:
                return {
                    "better_document": "Document 1",
                    "reason": f"Fewer risk factors ({risk_count1} vs {risk_count2})",
                    "recommendation": "Document 1 has fewer identified risks"
                }
            elif risk_count2 < risk_count1:
                return {
                    "better_document": "Document 2",
                    "reason": f"Fewer risk factors ({risk_count2} vs {risk_count1})",
                    "recommendation": "Document 2 has fewer identified risks"
                }
            else:
                return {
                    "better_document": "Similar",
                    "reason": "Both documents have comparable scores and risk profiles",
                    "recommendation": "Review both documents carefully; consider combining favorable terms"
                }


    def _identify_critical_differences(self, clause_diff: dict, entity_diff: dict, 
                                    rec_comp: dict) -> List[dict]:
        """Identify and prioritize critical differences between documents."""
        critical = []
        
        # Check for missing critical clauses
        critical_clauses = ["indemnification", "liability", "termination", 
                        "governing_law", "dispute_resolution"]
        for clause in critical_clauses:
            if clause in clause_diff:
                if clause_diff[clause]["status"] != "present_in_both":
                    critical.append({
                        "type": "missing_clause",
                        "severity": "HIGH",
                        "clause": clause,
                        "details": clause_diff[clause]["analysis"]
                    })
        
        # Check for party differences
        if "CONTRACT_PARTIES" in entity_diff:
            party_data = entity_diff["CONTRACT_PARTIES"]
            if party_data["only_in_doc1"] or party_data["only_in_doc2"]:
                critical.append({
                    "type": "party_difference",
                    "severity": "HIGH",
                    "details": f"Different parties identified: {party_data['only_in_doc1']} vs {party_data['only_in_doc2']}"
                })
        
        # Check for money amount differences
        if "MONEY_AMOUNTS" in entity_diff:
            money_data = entity_diff["MONEY_AMOUNTS"]
            if money_data["only_in_doc1"] or money_data["only_in_doc2"]:
                critical.append({
                    "type": "financial_difference",
                    "severity": "HIGH",
                    "details": f"Different monetary terms: {money_data['only_in_doc1']} vs {money_data['only_in_doc2']}"
                })
        
        # Check for significant score differences
        score_diff = abs(rec_comp["score_difference"])
        if score_diff > 15:
            critical.append({
                "type": "risk_score_difference",
                "severity": "MEDIUM",
                "details": f"Significant risk score difference: {score_diff:.1f} points"
            })
        
        return critical


    def _calculate_overall_similarity(self, clause_diff: dict, entity_diff: dict, 
                                    summary_similarity: float) -> dict:
        """Calculate overall document similarity score."""
        # Average clause similarity
        clause_sims = [c["similarity"] for c in clause_diff.values() if "similarity" in c]
        avg_clause_sim = sum(clause_sims) / len(clause_sims) if clause_sims else 0
        
        # Average entity similarity
        entity_sims = [e["similarity"] for e in entity_diff.values()]
        avg_entity_sim = sum(entity_sims) / len(entity_sims) if entity_sims else 0
        
        # Weighted overall similarity
        overall = (avg_clause_sim * 0.4 + avg_entity_sim * 0.3 + summary_similarity * 100 * 0.3)
        
        if overall >= 80:
            interpretation = "Documents are substantially similar"
        elif overall >= 60:
            interpretation = "Documents have moderate similarities with notable differences"
        elif overall >= 40:
            interpretation = "Documents have significant differences"
        else:
            interpretation = "Documents are substantially different"
        
        return {
            "percentage": round(overall, 2),
            "interpretation": interpretation,
            "clause_similarity": round(avg_clause_sim, 2),
            "entity_similarity": round(avg_entity_sim, 2),
            "summary_similarity": round(summary_similarity * 100, 2)
        }


   
        
from docx import Document

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Optional
import os

class ContractTemplateGenerator:
    """
    Enhanced contract generator with multiple template types.
    Supports filling placeholders with user data and generating downloadable contracts.
    """
    
    def __init__(self):
        self.templates = {
            "nda": self._get_nda_template(),
            "service_agreement": self._get_service_agreement_template(),
        }
    
    def _get_nda_template(self) -> str:
        """Non-Disclosure Agreement (NDA) Template"""
        return """
NON-DISCLOSURE AGREEMENT

This Non-Disclosure Agreement ("Agreement") is entered into as of {{effective_date}}, by and between:

DISCLOSING PARTY:
{{disclosing_party_name}}
{{disclosing_party_address}}
{{disclosing_party_email}}

and

RECEIVING PARTY:
{{receiving_party_name}}
{{receiving_party_address}}
{{receiving_party_email}}

(collectively referred to as the "Parties")

WHEREAS, the Disclosing Party possesses certain confidential and proprietary information related to {{business_purpose}};

WHEREAS, the Receiving Party desires to receive such confidential information for the purpose of {{purpose_of_disclosure}};

NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, the Parties agree as follows:

1. DEFINITION OF CONFIDENTIAL INFORMATION

"Confidential Information" means any and all information disclosed by the Disclosing Party to the Receiving Party, whether orally, in writing, or in any other form, including but not limited to:
   a) Technical data, trade secrets, know-how, research, product plans, products, services
   b) Software, algorithms, source code, object code
   c) Customer lists, supplier information, business strategies
   d) Financial information, forecasts, budgets
   e) Marketing plans, sales data
   f) Any other information marked as "Confidential" or that would reasonably be considered confidential

2. OBLIGATIONS OF RECEIVING PARTY

The Receiving Party agrees to:
   a) Hold and maintain the Confidential Information in strict confidence
   b) Not disclose the Confidential Information to any third parties without prior written consent
   c) Use the Confidential Information solely for the purpose of {{purpose_of_disclosure}}
   d) Protect the Confidential Information with the same degree of care used to protect its own confidential information, but in no case less than reasonable care
   e) Limit access to the Confidential Information to employees and contractors who have a legitimate need to know

3. EXCLUSIONS FROM CONFIDENTIAL INFORMATION

Confidential Information shall not include information that:
   a) Is or becomes publicly available through no breach of this Agreement
   b) Was rightfully in the Receiving Party's possession prior to disclosure
   c) Is independently developed by the Receiving Party without use of the Confidential Information
   d) Is rightfully received from a third party without breach of any confidentiality obligation

4. TERM AND TERMINATION

This Agreement shall commence on the Effective Date and continue for a period of {{term_years}} years, unless terminated earlier by either Party with {{notice_period}} days written notice. The obligations of confidentiality shall survive termination for a period of {{survival_years}} years.

5. RETURN OF MATERIALS

Upon termination of this Agreement or upon request by the Disclosing Party, the Receiving Party shall promptly return or destroy all Confidential Information and certify such destruction in writing.

6. NO LICENSE

Nothing in this Agreement grants the Receiving Party any license or right to the Confidential Information except as expressly stated herein.

7. REMEDIES

The Receiving Party acknowledges that breach of this Agreement may cause irreparable harm for which monetary damages would be inadequate. The Disclosing Party shall be entitled to seek equitable relief, including injunction and specific performance, in addition to all other remedies available at law or in equity.

8. GOVERNING LAW

This Agreement shall be governed by and construed in accordance with the laws of {{governing_state}}, without regard to its conflict of law provisions.

9. ENTIRE AGREEMENT

This Agreement constitutes the entire agreement between the Parties concerning the subject matter hereof and supersedes all prior agreements and understandings.

10. AMENDMENTS

This Agreement may only be amended or modified by a written document signed by both Parties.

IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

DISCLOSING PARTY:

_________________________________
{{disclosing_party_name}}
{{disclosing_party_title}}
Date: {{signature_date}}


RECEIVING PARTY:

_________________________________
{{receiving_party_name}}
{{receiving_party_title}}
Date: {{signature_date}}
"""

    def _get_service_agreement_template(self) -> str:
        """Service Agreement Template"""
        return """
SERVICE AGREEMENT

This Service Agreement ("Agreement") is entered into as of {{effective_date}}, by and between:

SERVICE PROVIDER:
{{provider_name}}
{{provider_address}}
{{provider_email}}
{{provider_phone}}

and

CLIENT:
{{client_name}}
{{client_address}}
{{client_email}}
{{client_phone}}

(collectively referred to as the "Parties")

WHEREAS, the Service Provider is engaged in the business of providing {{service_type}} services;

WHEREAS, the Client desires to engage the Service Provider to perform certain services as described herein;

NOW, THEREFORE, in consideration of the mutual covenants and agreements contained herein, the Parties agree as follows:

1. SERVICES

The Service Provider agrees to provide the following services ("Services"):

{{service_description}}

The Services shall be performed in accordance with industry standards and professional practices.

2. TERM

This Agreement shall commence on {{start_date}} and continue until {{end_date}} ("Initial Term"), unless terminated earlier in accordance with this Agreement. 

This Agreement may be renewed for additional {{renewal_period}} periods upon mutual written agreement of the Parties at least {{renewal_notice_days}} days prior to expiration.

3. COMPENSATION

3.1 Service Fees
The Client shall pay the Service Provider as follows:

Payment Structure: {{payment_structure}}
Total Amount: {{total_amount}}
Payment Schedule: {{payment_schedule}}

3.2 Additional Expenses
The Service Provider shall be reimbursed for pre-approved, reasonable expenses incurred in the performance of Services, including:
   - Travel expenses
   - Materials and supplies
   - Third-party services

All expenses must be documented with receipts and approved by the Client in writing prior to incurrence.

3.3 Payment Terms
Invoices shall be submitted {{invoice_frequency}} and are due within {{payment_terms}} days of receipt. Late payments shall accrue interest at {{late_fee_rate}}% per month.

4. DELIVERABLES

The Service Provider shall deliver the following to the Client:

{{deliverables}}

All deliverables shall be delivered by {{delivery_date}}, subject to any agreed-upon extensions.

5. INDEPENDENT CONTRACTOR

The Service Provider is an independent contractor and not an employee of the Client. The Service Provider shall be responsible for all taxes, insurance, and other obligations related to their status as an independent contractor.

6. INTELLECTUAL PROPERTY

6.1 Client Materials
All materials, information, and intellectual property provided by the Client remain the property of the Client.

6.2 Work Product
{{ip_ownership_clause}}

6.3 Pre-Existing Materials
Any pre-existing materials, tools, or intellectual property owned by the Service Provider prior to this Agreement shall remain the property of the Service Provider.

7. CONFIDENTIALITY

Both Parties agree to maintain the confidentiality of any proprietary or confidential information disclosed during the term of this Agreement. This obligation shall survive termination for a period of {{confidentiality_period}} years.

8. WARRANTIES

The Service Provider warrants that:
   a) Services will be performed in a professional and workmanlike manner
   b) Services will comply with all applicable laws and regulations
   c) Service Provider has the right and authority to enter into this Agreement
   d) Services will not infringe upon any third-party intellectual property rights

9. LIMITATION OF LIABILITY

IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOST PROFITS, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGES.

The Service Provider's total liability under this Agreement shall not exceed {{liability_cap}}.

10. TERMINATION

10.1 Termination for Convenience
Either Party may terminate this Agreement with {{termination_notice}} days written notice.

10.2 Termination for Cause
Either Party may terminate this Agreement immediately upon written notice if the other Party:
   a) Materially breaches this Agreement and fails to cure within {{cure_period}} days
   b) Becomes insolvent or files for bankruptcy
   c) Engages in fraudulent or illegal conduct

10.3 Effect of Termination
Upon termination, the Client shall pay for all Services performed up to the termination date and any expenses incurred.

11. DISPUTE RESOLUTION

Any disputes arising under this Agreement shall be resolved through:
{{dispute_resolution_method}}

12. GENERAL PROVISIONS

12.1 Entire Agreement
This Agreement constitutes the entire agreement between the Parties and supersedes all prior agreements.

12.2 Amendments
This Agreement may only be amended in writing signed by both Parties.

12.3 Assignment
Neither Party may assign this Agreement without the prior written consent of the other Party.

12.4 Governing Law
This Agreement shall be governed by the laws of {{governing_state}}.

12.5 Severability
If any provision of this Agreement is found invalid or unenforceable, the remaining provisions shall continue in full force and effect.

12.6 Notices
All notices under this Agreement shall be in writing and sent to the addresses listed above.

IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

SERVICE PROVIDER:

_________________________________
{{provider_name}}
{{provider_title}}
Date: {{signature_date}}


CLIENT:

_________________________________
{{client_name}}
{{client_title}}
Date: {{signature_date}}
"""

    def generate_contract(self, 
                         contract_type: str, 
                         user_data: Dict[str, str], 
                         output_path: str,
                         format_type: str = "docx") -> str:
        """
        Generate a contract from template with user data.
        
        Args:
            contract_type: Type of contract ("nda" or "service_agreement")
            user_data: Dictionary with placeholder values
            output_path: Path where the contract will be saved
            format_type: Output format ("docx" or "txt")
            
        Returns:
            Path to the generated contract file
        """
        if contract_type not in self.templates:
            raise ValueError(f"Unknown contract type: {contract_type}. Available: {list(self.templates.keys())}")
        
        # Get template
        template_content = self.templates[contract_type]
        
        # Add default values for common fields if not provided
        default_values = {
            "signature_date": datetime.now().strftime("%B %d, %Y"),
            "effective_date": datetime.now().strftime("%B %d, %Y"),
        }
        
        # Merge defaults with user data (user data takes precedence)
        filled_data = {**default_values, **user_data}
        
        # Fill in placeholders
        contract_content = template_content
        for key, value in filled_data.items():
            placeholder = f"{{{{{key}}}}}"
            contract_content = contract_content.replace(placeholder, str(value))
        
        # Check for unfilled placeholders
        import re
        unfilled = re.findall(r'\{\{([^}]+)\}\}', contract_content)
        if unfilled:
            print(f"Warning: The following placeholders were not filled: {', '.join(unfilled)}")
        
        # Generate output file
        if format_type == "docx":
            return self._generate_docx(contract_content, output_path, contract_type)
        else:
            return self._generate_txt(contract_content, output_path)
    
    def _generate_docx(self, content: str, output_path: str, contract_type: str) -> str:
        """Generate a formatted DOCX file"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add content with formatting
        lines = content.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            
            # Format titles (ALL CAPS lines)
            if line.isupper() and len(line) > 3:
                run.bold = True
                run.font.size = Pt(14)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Format section headers (lines starting with numbers)
            elif line[0].isdigit() and '.' in line[:3]:
                run.bold = True
                run.font.size = Pt(12)
            else:
                run.font.size = Pt(11)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _generate_txt(self, content: str, output_path: str) -> str:
        """Generate a plain text file"""
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        return output_path
    
    def get_required_fields(self, contract_type: str) -> list:
        """Get list of required fields for a contract type"""
        if contract_type not in self.templates:
            raise ValueError(f"Unknown contract type: {contract_type}")
        
        import re
        template = self.templates[contract_type]
        placeholders = re.findall(r'\{\{([^}]+)\}\}', template)
        
        # Remove duplicates and sort
        return sorted(list(set(placeholders)))
    
    def get_available_templates(self) -> Dict[str, str]:
        """Get list of available contract templates with descriptions"""
        return {
            "nda": "Non-Disclosure Agreement (NDA) - Protects confidential information between parties",
            "service_agreement": "Service Agreement - Defines terms for professional services between provider and client"
        }


